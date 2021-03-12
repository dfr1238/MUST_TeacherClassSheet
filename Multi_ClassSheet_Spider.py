from PySimpleGUI.PySimpleGUI import WIN_CLOSED
import selenium
from selenium import webdriver
import PySimpleGUI as sg
from docx import Document
from docx.shared import  Cm
from selenium.webdriver.support.ui import Select
import sys, os
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
import selenium.webdriver.support.ui as ui

class spider_Gui:

    def set_Input_Ready_Window():
        input_Ready_Layout =[
            [sg.Text('請到新開啟的網頁中登入')]
        ]
        return sg.Window('準備爬取資料',input_Ready_Layout,finalize=True)
    
    def set_finish_Window(spider):
        finsih_Window_Layout =[
            [sg.Text(f'###已抓取完該課表資料，如需下一筆請關閉 Word 後到網頁再次登入###\n{spider.info}\n請到程式目錄尋找 Word 檔：{spider.doc_name}')]
        ]
        return sg.Window('完成爬取',finsih_Window_Layout,finalize=True,modal=True)

    def set_running_Window(spider):
        running_Window_Layout =[
            [sg.Text(f'已偵測到可抓取表格！\n{spider.info}')]
        ]
        return sg.Window('已偵測到爬取資料！',running_Window_Layout,finalize=True)
    pass

class classMenu_Spider:
    table_Type=''
    doc_name=''
    table_Element =''
    table_TrList=''
    teacher_Name=''
    select_row_1=''
    select_row_2=''
    select_row_3=''
    driver=''
    info=''
    url='https://sss.must.edu.tw/RWD_CosInfo/service.asp#teachertab'

    def __init__(self) -> None:
        chrome_options = webdriver.ChromeOptions()
        #chrome_options.add_argument('--headless')
        chrome_options.add_argument('--disable-gpu')
        if __name__ == "__main__":

            if getattr(sys, 'frozen', False): 
                chrome_driver_path = os.path.join(sys._MEIPASS, 'chromedriver.exe')
                print(chrome_driver_path)
                self.driver = webdriver.Chrome(executable_path=chrome_driver_path,options=chrome_options)
            else:
                self.driver = webdriver.Chrome(options=chrome_options)
        try:
            self.driver.get(self.url)
            self.driver.maximize_window()
            self.driver.set_page_load_timeout(10)
        except selenium.common.exceptions.WebDriverException or selenium.common.TimeoutException:
            sg.popup_error(f'建立網頁驅動器時發生問題！請檢查網路連線與網頁 {self.url} 的狀態！')
            os._exit(0)

    def creating_Word(self): #建立 Word 檔
        self.driver.minimize_window()
        self.doc = Document()
        style = self.doc.styles['Normal']
        font = style.font
        font.size = Pt(12)
        p = self.doc.add_paragraph(self.info)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = self.doc.styles['Normal']
        table = self.doc.add_table(rows=1,cols=8)
        table.style = 'Light Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text ='節次'
        hdr_cells[1].text ='星期一'
        hdr_cells[2].text ='星期二'
        hdr_cells[3].text ='星期三'
        hdr_cells[4].text ='星期四'
        hdr_cells[5].text ='星期五'
        hdr_cells[6].text ='星期六'
        hdr_cells[7].text ='星期日'
        hdr_cells[0].width = Cm(.5)
        for i in range(1,8):
            hdr_cells[i].width = Cm(3)
        #date_heading=[['時段'],['星期一'],['星期二'],['星期三'],['星期四'],['星期五'],['星期六'],['星期日']]
        #print(date_heading)
        first_row=True
        for row in self.table_TrList:
            if(first_row):
                first_row=False
                continue
            tdlist = row.find_elements_by_tag_name('td')
            sg.one_line_progress_meter(f'匯出成 Word 檔中...',self.table_TrList.index(row),len(self.table_TrList)-1,'Progress',self.info)
            row_cells = table.add_row().cells
            td_count=0
            for td in tdlist:
                temp_text = td.text.split("\n")
                row_cells[td_count].text = str("\n".join(temp_text))
                td_count+=1

            #print('\n')
        sections = self.doc.sections
        for section in sections: #調整邊界
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)
        self.doc.save(self.doc_name)
        pass
    
    def classroom_table(self):
        try:
            wait = ui.WebDriverWait(self.driver,0.1)
            wait.until(lambda driver: driver.find_element_by_xpath('/html/body/div[1]/div/div/div/div/div[5]/div/center/div/table'))
            self.table_Element = self.driver.find_element_by_xpath('/html/body/div[1]/div/div/div/div/div[5]/div/center/div/table')
            self.table_TrList = self.driver.find_elements_by_tag_name('tr')
            self.table_Type = '教室課表'
            self.info = self.driver.find_element_by_xpath('/html/body/div[1]/div/div/div/div/div[5]/div/center/div/table/caption').text.split("\n")
            self.info = "\n".join(self.info)
            self.info = self.info.replace('搜尋條件 - ','')
            select_row_1= Select(self.driver.find_element_by_id('year_r'))
            select_row_2 = Select(self.driver.find_element_by_id('Divi_r'))
            select_row_3 = self.driver.find_element_by_xpath('/html/body/div[1]/div/div/div/div/div[5]/form/div[3]/div/button/div/div/div').text
            self.select_row_1 = select_row_1.first_selected_option.text
            self.select_row_2 = select_row_2.first_selected_option.text
            self.doc_name =(f'.\{self.select_row_1} - {self.select_row_2} - {select_row_3} - 課表.docx')
            return True
        except selenium.common.exceptions.TimeoutException:
            print('尚未找到教室課表元素！')
            return False
        pass

    def class_table(self):
        try:
            wait = ui.WebDriverWait(self.driver,0.1)
            wait.until(lambda driver: driver.find_element_by_xpath('/html/body/div[1]/div/div/div/div/div[4]/div/center/table[2]'))
            self.table_Element = self.driver.find_element_by_xpath('/html/body/div[1]/div/div/div/div/div[4]/div/center/table[2]')
            self.table_TrList = self.driver.find_elements_by_tag_name('tr')
            self.table_Type = '班級課表'
            self.info = self.driver.find_element_by_xpath('/html/body/div[1]/div/div/div/div/div[4]/div/center/table[2]/caption').text.split("\n")
            self.info = "\n".join(self.info)
            self.info = self.info.replace('搜尋條件 - ','')
            select_row_1= Select(self.driver.find_element_by_id('year_c'))
            select_row_2 = Select(self.driver.find_element_by_id('Divi_c'))
            select_row_3 = Select(self.driver.find_element_by_xpath('/html/body/div[1]/div/div/div/div/div[4]/form/div[3]/select'))
            select_row_4 = self.driver.find_element_by_xpath('/html/body/div[1]/div/div/div/div/div[4]/form/div[4]/div/button/div/div/div').text
            self.select_row_1 = select_row_1.first_selected_option.text
            self.select_row_2 = select_row_2.first_selected_option.text
            self.select_row_3 = select_row_3.first_selected_option.text
            self.doc_name =(f'.\{self.select_row_1} - {self.select_row_2} - {self.select_row_3} - {select_row_4} - 課表.docx')
            return True
        except selenium.common.exceptions.TimeoutException:
            print('尚未找到班級課表元素！')
            return False

    def teacher_table(self):
        try:
            wait = ui.WebDriverWait(self.driver,0.1)
            wait.until(lambda driver: driver.find_element_by_xpath('/html/body/div[1]/div/div/div/div/div[6]/div/center/div/table'))
            self.table_Element = self.driver.find_element_by_xpath('/html/body/div[1]/div/div/div/div/div[6]/div/center/div/table')
            self.table_TrList = self.driver.find_elements_by_tag_name('tr')
            self.table_Type = '教師課表'
            self.info = self.driver.find_element_by_xpath('/html/body/div[1]/div/div/div/div/div[6]/div/center/div/div').text.split("\n")
            self.info = "\n".join(self.info)
            self.info = self.info.replace('搜尋條件 - ','')
            select_row_1= Select(self.driver.find_element_by_id('ysList'))
            select_row_2 = Select(self.driver.find_element_by_id('Divi_t'))
            self.select_row_3 = self.driver.find_element_by_id('itea').get_attribute('value')
            self.select_row_1 = select_row_1.first_selected_option.text
            self.select_row_2 = select_row_2.first_selected_option.text
            self.doc_name =(f'.\{self.select_row_1} - {self.select_row_2} - {self.select_row_3} - 課表.docx')
            return True
        except selenium.common.exceptions.TimeoutException:
            print('尚未找到教師課表元素！')
            return False
    def check_table(self):
        try:
            if(self.teacher_table()):
                return True
            if(self.class_table()):
                return True
            if(self.classroom_table()):
                return True
            return False
        except selenium.common.exceptions.UnexpectedAlertPresentException:
            return False
            pass
        except selenium.common.exceptions.WebDriverException:
            sys.exit()

    def showing_data(self):
        sg.popup_notify(f'{self.info}',title='已找到課表！',display_duration_in_ms=150,fade_in_duration=150)
        pass

    def waiting_Input(self):
        ready_Window=None
        ready_Window=spider_Gui.set_Input_Ready_Window()
        running_Window=None
        finish_Window=None
        while True:
            window , event , values = sg.read_all_windows(500)
            state=self.check_table()
            if window == finish_Window:
                if event == WIN_CLOSED:
                    self.driver.quit()
                    window.close()
                    break
            if window == ready_Window:
                if event == WIN_CLOSED:
                    self.driver.quit()
                    window.close()
                    break
            if state:
                if(finish_Window!=None):
                    finish_Window.close()
                self.showing_data()
                ready_Window.close()
                running_Window=spider_Gui.set_running_Window(self)
                self.creating_Word()
                running_Window.close()
                self.driver.get(self.url)
                self.driver.maximize_window()
                docx=(self.doc_name)
                os.startfile(docx)
                finish_Window=spider_Gui.set_finish_Window(self)

Spider = classMenu_Spider()
Spider.waiting_Input()